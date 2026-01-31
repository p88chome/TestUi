# -*- coding: utf-8 -*-
from docx import Document
import os

base_path = r"c:\Users\user\OneDrive\桌面\All平台\文件測試"

# 讀取管理辦法
policy_path = os.path.join(base_path, "採購驗收管理辦法.docx")
doc = Document(policy_path)
print("=" * 60)
print("=== 採購驗收管理辦法 ===")
print("=" * 60)
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)

# 讀取訪談紀錄
for i in range(1, 4):
    interview_path = os.path.join(base_path, f"[永瑞] 採購驗收管理辦法訪談 ({i}).docx")
    if os.path.exists(interview_path):
        doc = Document(interview_path)
        print("\n" + "=" * 60)
        print(f"=== 訪談紀錄 ({i}) ===")
        print("=" * 60)
        for p in doc.paragraphs:
            if p.text.strip():
                print(p.text)
