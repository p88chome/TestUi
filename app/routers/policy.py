"""
Policy Gap Analysis API Router
提供制度差異分析的專屬 API
"""

import io
import os
import re
import uuid
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse
import json
import asyncio
from typing import List, Optional
from sqlalchemy.orm import Session

from app.api.deps import get_db, get_current_user
from app.models.user import User

router = APIRouter(prefix="/policy", tags=["policy"])

ACCEPTED_EXTS = ('.docx', '.pdf')


def _validate_file(file: UploadFile) -> None:
    if not file.filename:
        raise HTTPException(status_code=400, detail="檔案名稱無效")
    if not file.filename.lower().endswith(ACCEPTED_EXTS):
        raise HTTPException(
            status_code=400,
            detail=f"不支援的檔案格式：{file.filename}。支援 .docx 和 .pdf"
        )


def _extract_text_from_bytes(filename: str, content: bytes) -> str:
    """從 DOCX 或 PDF 位元組中擷取純文字。"""
    name_lower = filename.lower()
    if name_lower.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"無法解析 DOCX：{e}")
    if name_lower.endswith('.pdf'):
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return "\n".join(parts)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"無法解析 PDF：{e}")
    return content.decode('utf-8', errors='ignore')


@router.post("/analyze")
async def analyze_policy_gap(
    policy_file: UploadFile = File(..., description="管理辦法文件 (docx)"),
    interview_files: List[UploadFile] = File(..., description="訪談紀錄文件列表 (docx)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    制度差異分析 API
    
    上傳管理辦法和訪談紀錄，AI 自動比對差異並產出 Markdown 修改建議報告。
    
    - **policy_file**: 管理辦法文件 (docx 格式)
    - **interview_files**: 訪談紀錄文件列表 (docx 格式，可多份)
    
    Returns:
        Markdown 格式的差異分析報告
    """
    
    # 驗證檔案類型（DOCX / PDF）
    _validate_file(policy_file)
    for f in interview_files:
        _validate_file(f)
    
    async def generate_analyze_stream():
        try:
            # 1. Parsing in memory
            policy_raw = await policy_file.read()
            policy_content = _extract_text_from_bytes(policy_file.filename, policy_raw)
            
            yield json.dumps({"status": "extracting", "message": "已讀取管理辦法..."}) + "\n"
            
            interview_contents = []
            for i_f in interview_files:
                i_raw = await i_f.read()
                text = _extract_text_from_bytes(i_f.filename, i_raw)
                interview_contents.append(f"【訪談紀錄: {i_f.filename}】\n{text}")
            
            yield json.dumps({"status": "extracting", "message": f"已讀取 {len(interview_files)} 份訪談紀錄，開始分析..."}) + "\n"
                
            combined_interviews = "\n\n---\n\n".join(interview_contents)
            
            from app.core.config import settings
            import httpx
            
            # --- Prompt Logic ---
            system_prompt = """你是一位資深的內控顧問與稽核專家。

你的任務是比對「管理辦法」與「訪談紀錄」之間的差異，找出：
1. 制度有規定但實務未落實的地方
2. 實務有執行但制度未明文規定的地方
3. 制度與實務不一致之處
4. 建議新增或修改的條款

## 分析原則
- 精確對應：每個差異必須指出管理辦法的具體條款位置
- 客觀中立：如實呈現差異，不做主觀臆測
- 具體可行：建議修改必須具體，可直接採用

## 輸出格式 (Markdown)

# 制度差異分析報告

**分析日期**: [今天日期]
**管理辦法**: [文件名稱]

---

## 一、管理辦法摘要

(列出管理辦法的主要條款結構，每條簡要說明)

---

## 二、差異分析

| 條款位置 | 管理辦法規定 | 訪談實務發現 | 差異類型 | 建議修改 |
|---------|-------------|-------------|---------|---------|

### 詳細說明

(對每個重要差異進行詳細說明)

---

## 三、建議新增條款

(訪談中提到但管理辦法沒有的控制點，建議新增)

---

## 四、總結與優先順序

| 優先級 | 修改項目 | 理由 |
|-------|---------|------|

---

請用繁體中文輸出，保持專業客觀的語氣。"""

            user_prompt = f"""請根據以下資料進行制度差異分析：

## 管理辦法
**文件名稱**: {policy_file.filename}

{policy_content}

---

## 訪談紀錄

{combined_interviews}

---

請產出完整的差異分析報告。"""

            from app.models.domain import AIModel
            from app.models.stats import UsageLog
            active_model = db.query(AIModel).filter(AIModel.is_active == True).first()
            if active_model and active_model.deployment_name:
                deployment = active_model.deployment_name
                model_name_for_log = active_model.name
            else:
                deployment = getattr(settings, "AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4.1")
                model_name_for_log = "Azure OpenAI GPT-4.1"

            api_key = settings.AZURE_OPENAI_API_KEY
            endpoint = settings.AZURE_OPENAI_ENDPOINT
            api_version = getattr(settings, "AZURE_OPENAI_API_VERSION", "2024-12-01-preview")
            
            if not api_key:
                yield json.dumps({"status": "error", "error": "未設定 Azure OpenAI API Key"}) + "\n"
                return
                
            url = f"{endpoint}openai/deployments/{deployment}/chat/completions?api-version={api_version}"
            headers = {"Content-Type": "application/json", "api-key": api_key}
            payload = {
                "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                "temperature": 0.2, 
                "stream": True,
                "stream_options": {"include_usage": True}
            }
            
            async with httpx.AsyncClient(timeout=300.0) as client:
                async with client.stream("POST", url, headers=headers, json=payload) as response:
                    if response.status_code != 200:
                        err_text = await response.aread()
                        yield json.dumps({"status": "error", "error": f"AI 服務傳回錯誤：{err_text.decode()}"}) + "\n"
                        return

                    async for line in response.aiter_lines():
                        if not line or not line.startswith("data: "): continue
                        if line.strip() == "data: [DONE]": break
                        
                        try:
                            data = json.loads(line[6:])
                            
                            # 擷取 OpenAI 串流最後一包的 Token 用量
                            if "usage" in data and data["usage"]:
                                usage = data["usage"]
                                total_tokens = usage.get("total_tokens", 0)
                                if total_tokens > 0:
                                    try:
                                        log = UsageLog(
                                            user_id=current_user.id,
                                            app_name="Policy Analysis",
                                            model_name=model_name_for_log,
                                            total_tokens=total_tokens,
                                            estimated_cost=total_tokens * 0.000003  # 粗略估算成本
                                        )
                                        db.add(log)
                                        db.commit()
                                    except Exception as db_err:
                                        print(f"Failed to log usage: {db_err}")

                            # 一般內容串流 (最後一包含有 usage 的 chunk 其 choices 會是空的)
                            if data.get('choices'):
                                delta = data['choices'][0].get('delta', {})
                                content = delta.get('content', '')
                                if content:
                                    yield json.dumps({"status": "content", "content": content}) + "\n"
                        except:
                            continue
            
            yield json.dumps({"status": "done"}) + "\n"
            
        except Exception as e:
            yield json.dumps({"status": "error", "error": f"分析過程發生錯誤：{str(e)}"}) + "\n"

    return StreamingResponse(generate_analyze_stream(), media_type="application/x-ndjson")


@router.post("/analyze-text")
async def analyze_policy_gap_text(
    policy_content: str = Form(..., description="管理辦法文字內容"),
    interview_content: str = Form(..., description="訪談紀錄文字內容"),
    policy_name: str = Form("管理辦法", description="管理辦法名稱"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    制度差異分析 API (純文字版本) - 支援 Streaming
    """
    async def generate_text_stream():
        try:
            yield json.dumps({"status": "processing", "message": "收到文字內容，開始分析..."}) + "\n"
            
            from app.core.config import settings
            import httpx
            
            system_prompt = """你是一位資深的內控顧問與稽核專家。比對「管理辦法」與「訪談紀錄」之間的差異，找出制度落差並給予修改建議。請產出包含摘要、差異表格、詳細說明與總結的專業 Markdown 報告。請用繁體中文輸出。"""
            user_prompt = f"""請根據以下資料進行制度差異分析：

## 管理辦法
**文件名稱**: {policy_name}

{policy_content}

---

## 訪談紀錄

{interview_content}

---

請產出完整的差異分析報告。"""

            from app.models.domain import AIModel
            from app.models.stats import UsageLog
            active_model = db.query(AIModel).filter(AIModel.is_active == True).first()
            if active_model and active_model.deployment_name:
                deployment = active_model.deployment_name
                model_name_for_log = active_model.name
            else:
                deployment = getattr(settings, "AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4.1")
                model_name_for_log = "Azure OpenAI GPT-4.1"

            api_key = settings.AZURE_OPENAI_API_KEY
            endpoint = settings.AZURE_OPENAI_ENDPOINT
            api_version = getattr(settings, "AZURE_OPENAI_API_VERSION", "2024-12-01-preview")
            
            url = f"{endpoint}openai/deployments/{deployment}/chat/completions?api-version={api_version}"
            headers = {"Content-Type": "application/json", "api-key": api_key}
            payload = {
                "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                "temperature": 0.2, 
                "stream": True,
                "stream_options": {"include_usage": True}
            }
            
            async with httpx.AsyncClient(timeout=300.0) as client:
                async with client.stream("POST", url, headers=headers, json=payload) as response:
                    if response.status_code != 200:
                        yield json.dumps({"status": "error", "error": f"AI 服務異常 (HTTP {response.status_code})"}) + "\n"
                        return

                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            if line.strip() == "data: [DONE]": break
                            try:
                                data = json.loads(line[6:])
                                
                                if "usage" in data and data["usage"]:
                                    usage = data["usage"]
                                    total_tokens = usage.get("total_tokens", 0)
                                    if total_tokens > 0:
                                        try:
                                            log = UsageLog(
                                                user_id=current_user.id,
                                                app_name="Policy Analysis (Text)",
                                                model_name=model_name_for_log,
                                                total_tokens=total_tokens,
                                                estimated_cost=total_tokens * 0.000003
                                            )
                                            db.add(log)
                                            db.commit()
                                        except Exception as db_err:
                                            print(f"Failed to log usage: {db_err}")

                                if data.get('choices'):
                                    content = data['choices'][0].get('delta', {}).get('content', '')
                                    if content:
                                        yield json.dumps({"status": "content", "content": content}) + "\n"
                            except: continue
            
            yield json.dumps({"status": "done"}) + "\n"
        except Exception as e:
            yield json.dumps({"status": "error", "error": str(e)}) + "\n"

    return StreamingResponse(generate_text_stream(), media_type="application/x-ndjson")

# ─── Word Export ─────────────────────────────────────────────────────────────

from pydantic import BaseModel

class ExportDocxRequest(BaseModel):
    report: str               # The full Markdown report text
    policy_name: str = "制度差異分析報告"


@router.post("/export-docx")
async def export_policy_docx(
    request: ExportDocxRequest,
    current_user: User = Depends(get_current_user),
):
    """
    將 Markdown 格式的差異分析報告轉換為可下載的 Word (.docx) 文件。
    使用 python-docx 產生帶有標題、表格、段落的完整 Word 文件。
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        doc = DocxDocument()

        # ── Page setup ──
        section = doc.sections[0]
        section.page_width  = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin  = section.bottom_margin = Inches(1.0)

        # ── Map Markdown to Word ──
        if request.report:
            lines = request.report.split('\n')
        else:
            lines = ["沒有分析報告內容"]

        i = 0
        while i < len(lines):
            line = lines[i]

            # H1
            if line.startswith('# ') and not line.startswith('## '):
                h = doc.add_heading(line[2:].strip(), level=1)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x00, 0x5B, 0x2D)

            # H2
            elif line.startswith('## ') and not line.startswith('### '):
                h = doc.add_heading(line[3:].strip(), level=2)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x00, 0x5B, 0x2D)

            # H3
            elif line.startswith('### '):
                h = doc.add_heading(line[4:].strip(), level=3)

            # Horizontal rule
            elif line.strip() == '---':
                doc.add_paragraph('─' * 40)

            # Table  (| col | col | …)
            elif line.strip().startswith('|') and '|' in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_raw = lines[i].strip()
                    if not re.match(r'^\|[\s\-|:]+\|$', row_raw):
                        cells = [c.strip() for c in row_raw.strip('|').split('|')]
                        table_lines.append(cells)
                    i += 1

                if table_lines:
                    ncols = max(len(r) for r in table_lines)
                    tbl = doc.add_table(rows=len(table_lines), cols=ncols)
                    tbl.style = 'Table Grid'
                    for ri, row_cells in enumerate(table_lines):
                        for ci, cell_text in enumerate(row_cells):
                            cell = tbl.cell(ri, ci)
                            cell.text = cell_text
                            if ri == 0:  # Header row bold
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                    doc.add_paragraph('')
                continue

            # Bullet list
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')

            # Numbered list
            elif re.match(r'^\d+\.\s', line.strip()):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')

            # Bold inline **text**
            elif line.strip():
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line.strip())
                for pi, part in enumerate(parts):
                    if not part: continue
                    run = p.add_run(part)
                    if pi % 2 == 1:   # odd index = bold content
                        run.bold = True
                if p.runs:
                    p.runs[0].font.size = Pt(10.5)
            else:
                doc.add_paragraph('')  # blank line → spacing

            i += 1

        # ── Stream response ──
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        safe_name = request.policy_name.replace('/', '_').replace('\\', '_')
        date_str = datetime.now().strftime('%Y%m%d')
        filename = f"{safe_name}_{date_str}.docx"

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
        )

    except Exception as e:
        # 回傳 400 Bad Request 錯誤，強迫前端收到正確 JSON 解析而非觸發 Azure 的 500 CORS 清洗
        raise HTTPException(
            status_code=400,
            detail=f"Word 匯出發生錯誤: {str(e)}"
        )

