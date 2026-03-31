"""
Meeting Minutes API Router
提供會議紀錄助手的專屬 API
"""

import os
import uuid
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List
from sqlalchemy.orm import Session

from app.api.deps import get_db, get_current_user
from app.models.user import User
from app.services.gateway import LLMGateway

router = APIRouter(prefix="/meeting", tags=["meeting"])


class MeetingGenerateRequest(BaseModel):
    """會議紀錄生成請求"""
    content: str
    meeting_title: Optional[str] = None
    meeting_date: Optional[str] = None
    participants: Optional[str] = None


class MeetingGenerateResponse(BaseModel):
    """會議紀錄生成回應"""
    status: str
    meeting_minutes: Optional[str] = None
    format: str = "markdown"
    meeting_title: Optional[str] = None
    meeting_date: Optional[str] = None
    error: Optional[str] = None


@router.post("/generate", response_model=MeetingGenerateResponse)
async def generate_meeting_minutes(
    request: MeetingGenerateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    從文字內容生成結構化會議紀錄
    
    - **content**: 會議對話或逐字稿 (必填)
    - **meeting_title**: 會議標題 (選填)
    - **meeting_date**: 會議日期 (選填，預設今天)
    - **participants**: 出席者 (選填)
    """
    
    if not request.content or not request.content.strip():
        raise HTTPException(status_code=400, detail="會議內容不可為空")
    
    # 準備參數
    meeting_date = request.meeting_date or datetime.now().strftime("%Y-%m-%d")
    
    # 建構 Prompt
    # 建構 Prompt (基於使用者建議的「管理顧問」Persona)
    system_prompt = """你是一位熟悉會計、稅務與財務結構的管理顧問，擅長將口語或零散筆記整理為正式的中文會議紀錄。

你的任務是將非結構化的會議對話或文字，轉換為一份正式、條理清楚的會議紀錄文件，適合提供給管理階層與查核人員閱讀。

## 處理原則
1. **角色定位**：以專業第三方或內控稽核視角進行紀錄。
2. **語氣要求**：採用正式、偏書面的中文語氣（例如使用「說明」、「指出」、「建議」、「需進一步確認」），避免口語化。
3. **資訊忠實**：
   - **絕對禁止**自行新增任何金額、年份或具體數字。
   - **絕對禁止**更改 OU 編號、公司名稱（如 Elan Hong Kong、里弗爾）、關係人名稱。
   - 必須完整保留關鍵數字與交易結構描述。
4. **結構重組**：
   - 不要只做逐字翻譯，請依照議題邏輯重新組織內容。
   - 若原文已有明確小節（如 3.1, 3.2），請保留並優化。

## 輸出結構參考
請依照內容邏輯，整理為類似以下的結構（視實際內容調整章節名稱）：

**[文件標題]** (若使用者未提供，請根據內容自動訂定)

**一、會議目的**
(簡述會議主要目標)

**二、重點議題與討論摘要**
(將內容依主題分段，例如「費用概況」、「財務成本」、「人事成本」等)
- 使用條列式 + 短段落。
- 每一點需自成完整句子。

**三、決議事項**
1. [明確的決議]
2. [包含具體數字或方向]

**四、待辦與後續行動**
| 負責人/單位 | 待辦事項 | 預計時程 |
|-------------|----------|----------|
| [人名] | [任務] | [日期] |

## 注意事項
- 若內容提及具體交易結構（如 A -> B -> C），描述必須完整。
- 若資訊不明確，請標記「(待確認)」，不要瞎掰。
- 輸出格式請使用標準 Markdown。"""

    user_prompt = f"""請根據以下會議內容，產出結構化會議紀錄：

{f'**會議標題**: {request.meeting_title}' if request.meeting_title else ''}
{f'**會議日期**: {meeting_date}' if meeting_date else ''}
{f'**出席者**: {request.participants}' if request.participants else ''}

---
**會議內容**：

{request.content}
"""

    try:
        gateway = LLMGateway(db)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        response = await gateway.chat(messages=messages, temperature=0.3)
        
        choices = response.get("choices", [])
        if choices:
            minutes_content = choices[0].get("message", {}).get("content", "")
            return MeetingGenerateResponse(
                status="success",
                meeting_minutes=minutes_content,
                format="markdown",
                meeting_title=request.meeting_title or "會議紀錄",
                meeting_date=meeting_date
            )
        else:
            return MeetingGenerateResponse(
                status="error",
                error="AI 未返回有效內容"
            )
            
    except Exception as e:
        return MeetingGenerateResponse(
            status="error",
            error=f"生成會議紀錄時發生錯誤: {str(e)}"
        )


@router.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    method: str = Form("whisper"),
    model_size: str = Form("base"),
    language: str = Form("zh"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    音訊/影片轉文字（非同步版本）
    
    上傳後立即回傳 task_id，轉錄在背景執行。
    使用 GET /meeting/task/{task_id} 查詢進度。
    
    - **file**: 音訊或影片檔案 (mp3, wav, m4a, mp4, webm)
    - **method**: 識別方法 (whisper 或 google)
    - **model_size**: Whisper 模型大小 (tiny, base, small, medium, large)
    - **language**: 語言代碼 (zh, en, zh-TW, zh-CN)
    """
    import tempfile
    import logging
    from app.services.task_manager import task_manager

    logger = logging.getLogger(__name__)
    
    # 驗證檔案類型
    allowed_types = [".mp3", ".wav", ".m4a", ".webm", ".ogg", ".mp4", ".avi", ".mov", ".mkv"]
    file_ext = os.path.splitext(file.filename)[1].lower() if file.filename else ""
    
    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"不支援的檔案格式。支援格式: {', '.join(allowed_types)}"
        )
    
    # 先將檔案存到暫存路徑
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            content = await file.read()
            tmp.write(content)
            temp_filepath = tmp.name
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"檔案儲存失敗: {str(e)}")

    # 定義背景轉錄 coroutine
    async def _do_transcribe():
        import asyncio
        try:
            from app.skills.transcription import run as transcribe_run
            # 在 executor 中執行同步的 transcribe（避免阻塞 event loop）
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                lambda: transcribe_run(
                    file_path=temp_filepath,
                    method=method,
                    model_size=model_size,
                    language=language,
                    include_timestamps=True
                )
            )
            if result['success']:
                return {
                    "status": "success",
                    "transcript": result['text'],
                    "language": result.get('language', language),
                    "segments": result.get('segments', []),
                    "filename": file.filename,
                    "method": method
                }
            else:
                raise Exception(f"轉錄失敗: {result.get('error', '未知錯誤')}")
        finally:
            # 確保暫存檔案被清理
            if os.path.exists(temp_filepath):
                try:
                    os.unlink(temp_filepath)
                except OSError:
                    pass

    # 提交背景任務
    task_id = await task_manager.submit(_do_transcribe())
    logger.info(f"Transcription task submitted: {task_id}, file: {file.filename}")

    return {
        "status": "processing",
        "task_id": task_id,
        "message": "音檔已收到，正在背景轉錄中。請使用 task_id 查詢進度。"
    }


@router.get("/task/{task_id}")
async def get_task_status(
    task_id: str,
    current_user: User = Depends(get_current_user)
):
    """
    查詢背景任務狀態（如 Whisper 轉錄）
    
    - **task_id**: 由 /transcribe 回傳的任務 ID
    
    回傳:
    - status: pending / running / completed / failed
    - result: 完成時的轉錄結果
    - error: 失敗時的錯誤訊息
    """
    from app.services.task_manager import task_manager

    status = task_manager.get_status(task_id)
    if status is None:
        raise HTTPException(status_code=404, detail="找不到該任務，可能已過期或不存在")
    
    return status


@router.post("/export")
async def export_meeting_minutes(
    content: str = Form(...),
    format: str = Form("markdown"),
    title: str = Form("meeting_minutes"),
    current_user: User = Depends(get_current_user)
):
    """
    匯出會議紀錄
    
    - **content**: 會議紀錄內容 (Markdown)
    - **format**: 匯出格式 (markdown, docx)
    - **title**: 檔案名稱
    """
    
    # 建立暫存目錄
    temp_dir = os.path.join(os.getcwd(), "temp")
    os.makedirs(temp_dir, exist_ok=True)
    
    safe_title = "".join(c for c in title if c.isalnum() or c in ('-', '_', ' ')).strip()
    if not safe_title:
        safe_title = "meeting_minutes"
    
    if format == "markdown":
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        filepath = os.path.join(temp_dir, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        
        return FileResponse(
            filepath,
            media_type="text/markdown",
            filename=filename
        )
    
    elif format == "docx":
        try:
            from docx import Document
            
            filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(temp_dir, filename)
            
            doc = Document()
            
            # 簡單解析 Markdown 並轉換為 Word
            lines = content.split("\n")
            for line in lines:
                if line.startswith("## "):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith("| "):
                    # 簡化處理表格 - 轉為段落
                    if not line.startswith("|---"):
                        doc.add_paragraph(line.replace("|", " | "))
                elif line.strip():
                    doc.add_paragraph(line)
            
            doc.save(filepath)
            
            return FileResponse(
                filepath,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=filename
            )
            
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="Word 匯出功能需要 python-docx 套件"
            )
    
    else:
        raise HTTPException(
            status_code=400,
            detail=f"不支援的匯出格式: {format}"
        )
