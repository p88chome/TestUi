from docx import Document
import os

def execute(input_data: dict) -> dict:
    """
    Executes Word Docx operations.
    """
    operation = input_data.get("operation")
    output_path = input_data.get("output_path", "report.docx")
    
    # Ensure directory
    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir)

    try:
        if operation == "create_report":
            content = input_data.get("content", [])
            return _create_report(output_path, content)
        elif operation == "append_table":
            headers = input_data.get("headers", [])
            rows = input_data.get("rows", [])
            return _append_table(output_path, headers, rows)
        else:
            return {"error": f"Unknown operation: {operation}"}
    except Exception as e:
        return {"error": str(e)}

def _create_report(output_path, content):
    doc = Document()
    
    for paragraph_text in content:
        text = str(paragraph_text)
        if text.startswith("# "):
            doc.add_heading(text[2:], level=1)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=2)
        elif text.startswith("### "):
            doc.add_heading(text[4:], level=3)
        else:
            doc.add_paragraph(text)
            
    doc.save(output_path)
    return {"status": "success", "message": f"Created report at {output_path}"}

def _append_table(output_path, headers, rows):
    if os.path.exists(output_path):
        doc = Document(output_path)
    else:
        doc = Document()

    # Add optional spacing
    doc.add_paragraph() 
    
    table = doc.add_table(rows=1, cols=len(headers))
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        
    # Add rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = str(item)
                
    doc.save(output_path)
    return {"status": "success", "message": f"Appended table to {output_path}"}
